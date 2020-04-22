import georinex as gr
import numpy as np
import numpy.linalg as lg
import math 
import matplotlib.pyplot as plt
import os
import shutil
import json
from functools import reduce
from docx import Document
from docx.oxml.ns import qn

def as_num3(x):
    return float('{:.3f}'.format(x))

# 不同观测站对比的精度
def as_num(x):
    return '{:.6f}'.format(x)

class pyec:

    #数据
    svs = []  #可以选择的卫星
    times = [] # 可以选的时间
    OmegaDote = 7.2921151467e-5 #地球自传参数

    def __init__(self,nav,sp3,sv_need=2,drawmode='auto',n=1,a=1,savepath=r'./数据输出'):  
        print()
        print()      
        print('当前处理{}批，共{}批'.format(n,a))
        print('导航文件：',nav)
        print('星历文件：',sp3)        
        self.nav = gr.load(nav)
        self.sp3 = gr.load(sp3)
        self.t0 = np.datetime64(self.sp3.t0) #datetime64属性
        self.sv_need = sv_need
        self.drawmode = drawmode
        if n==1:
            if not os.path.exists(savepath):
                os.mkdir(savepath)
            else:
                shutil.rmtree(savepath)
                os.mkdir(savepath)        
        navname = str(self.nav.filename).split('_')[0].split('00')[0]
        navtime = str(self.t0).split('T')[0]
        if not os.path.exists(savepath +r'/'+navname):
            os.mkdir(savepath+r'/' + navname)
        if not os.path.exists(savepath+r'/' + navname + r'/' + navtime):
            os.mkdir(savepath+r'/'+navname + r'/' + navtime)
        self.savepath = savepath+ r'/' + navname + r'/' + navtime + r'/'
        self.logfile = open(savepath+ r'/' + navname + r'/' + navtime + r'/' + r'/output.txt',mode='w',encoding='utf-8')
        self.logfile.write('导航文件:{}\r\n'.format(nav))
        self.logfile.write('星历文件:{}\r\n'.format(sp3))
        self.navname = navname


    def get_available_data(self,sv):
        nav = self.nav
        sp3 = self.sp3

        # 获取他们都有的时间
        # 时间线最大，卫星保底4科
        nav_times = set(nav.time.values)
        sp3_times = set(sp3.time.values)
        times_both = list(nav_times & sp3_times)
        timess = np.sort(times_both)
        if len(timess) <= 0:
            print('日期不同')
            return None

        #print(timess)

        # 根据都有的时间获取导航文件ok的卫星
        svs_sets = []
        for t in timess:
            #在t时间有的卫星
            toolman = set((nav.sel(time=t).dropna(dim='sv',how='all').sv.values))
            svs_sets.append(toolman)

        for i in range(len(svs_sets)):
            temp = svs_sets[i] & svs_sets[i+1]
            if  temp:
                svs_sets[i+1] = temp
                #至少4颗卫星
                if len(temp)<sv:
                    #迭代i次
                    '''
                    print('卫星保底模式。')
                    print('共有时间线:',i)
                    print('共有卫星：',svs_sets[i])
                    '''
                    self.svs = list(svs_sets[i])
                    self.times = timess[0:i]
                    break
                continue
            else:
                '''
                print('时间线最优化模式。')
                print('共有时间线',i)
                print('共有卫星:',svs_sets[i])
                '''
                self.times = timess[0:i]
                self.svs = list(svs_sets[i])
                break
        
        self.logfile.write('共有最大时间线:{}\r\n'.format(self.times))
        self.logfile.write('共有最大卫星:{}\r\n'.format(self.svs))
        self.logfile.write('\r\n\r\n')

    def calEk(self,Eq,M,es):
        E = M + es * np.sin(Eq)
        if np.fabs(E-Eq) < 1e-15:
            return E
        else:
            return self.calEk(E,M,es)

    def cal(self,t,data):
        OmegaDote = self.OmegaDote
        # 1.计算规化时间
        tk = t - data['Toe'].values
        #print('tk:',tk)
        # 计算卫星平均角速度
        sqrtGM = np.sqrt(3.986005e+14)
        sqrtA = data['sqrtA'].values
        n0 = np.divide(sqrtGM, pow(sqrtA,3))
        n = n0 + data['DeltaN'].values
        # 3. 信号发射时刻的平定点角
        Mk = data['M0'].values + n * tk
        if Mk > np.pi * 2 :
            Mk = Mk - np.pi * 2
        elif Mk< 0:
            Mk = Mk + np.pi *2
        # 4.偏近点角Ek
        es = data['Eccentricity'].values
        Ek = self.calEk(Mk,Mk,es)
        # 5.真近点角vk
        vk_top_cos = np.cos(Ek) - es 
        vk_top_sin = np.sqrt(1 - es*es) * np.sin(Ek)
        vk_bottom = 1 - es * np.cos(Ek)
        vk_cos = np.divide(vk_top_cos,vk_bottom)
        vk_sin = np.divide(vk_top_sin,vk_bottom)
        vk = math.atan2(vk_sin,vk_cos)
        # 6.升交点角距phik
        phik = vk + data['omega'].values
        # 7.摄动矫正项 Duk,Drk,Dik
        Duk = data['Cus'].values * np.sin(2*phik) + data['Cuc'].values * np.cos(2*phik)
        Drk = data['Crs'].values * np.sin(2*phik) + data['Crc'].values * np.cos(2*phik)
        Dik = data['Cis'].values * np.sin(2*phik) + data['Cic'].values * np.cos(2*phik)
        # 8.矫正后的升交点角距uk,矢径长度rk，轨道倾角ik
        uk = phik + Duk
        rk = sqrtA * sqrtA * (1 - es * np.cos(Ek)) + Drk
        ik = data['Io'].values+ data['IDOT'].values * tk + Dik
        # 9.计算卫星发射时刻轨道平面的位置(xk,yk)
        xk = rk * np.cos(uk)
        yk = rk * np.sin(uk)
        # 10计算信号发射时刻的升交点赤经Omegak
        Omegak = data['Omega0'].values + (data['OmegaDot'].values - OmegaDote) * tk - OmegaDote * data['Toe'].values
        # 11.计算卫星的坐标(x,y,z)
        x = xk * np.cos(Omegak) - yk * np.cos(ik) * np.sin(Omegak)
        y = xk * np.sin(Omegak) + yk * np.cos(ik) * np.cos(Omegak)
        z = yk * np.sin(ik)

        
        #print('计算坐标:',x,y,z)
        self.x = as_num3(x)
        self.y = as_num3(y)
        self.z = as_num3(z)
        #print('{}时的卫星坐标WGS84:({},{},{})'.format(data.time.values,x,y,z))


    def transTime(self):
        pass


    def cal_sd(self):
        '''
        计算均方差
        '''

    def draw(self,datas,mode,display):
        sv_num = len(datas)
        savepath = self.savepath

        def sub(savepath,sv_num,datas,display):
            plt.figure(1)
            #plt.rcParams.update({"font.size":20})
            plt.rcParams['font.sans-serif']=['Hei']
            plt.rcParams['axes.unicode_minus']=False
            i = 1           
            for prn,data in datas.items():
            
                plt.subplot(sv_num,1,i)                
                i = i+1
                plt.title(self.navname + '观测站  ' + prn+ '  ' + str(self.t0).split('T')[0]+'误差分析')
                plt.plot(data['time'],data['x'],'k-o',label='x坐标')
                plt.plot(data['time'],data['y'],'k-v',label='y坐标')
                plt.plot(data['time'],data['z'],'k-s',label='z坐标')
                #plt.plot(data['time'],np.zeros(len(data['x'])),'k--')
                plt.ylim((-2,2))
                plt.xlabel('时间')
                plt.ylabel('误差(m)')
                plt.tight_layout()
                plt.legend(loc=1)   
            plt.savefig(savepath + 'prn.png')   
            if display==True:         
                plt.show()
            plt.clf()
            
        
        def single(savepath,sv_num,datas,display):
            plt.rcParams['font.sans-serif']=['Hei']
            plt.rcParams['axes.unicode_minus']=False
            #plt.rcParams.update({"font.size":20})
            i = 1
            
            for prn,data in datas.items():
                plt.figure(i)
                i = i+1
                plt.title(self.navname + '观测站  ' + prn+ '  ' + str(self.t0).split('T')[0]+'误差分析')
                plt.plot(data['time'],data['x'],'k-o',label='x坐标')
                plt.plot(data['time'],data['y'],'k-v',label='y坐标')
                plt.plot(data['time'],data['z'],'k-s',label='z坐标')
                plt.ylim((-2,2))
                plt.xlabel('时间')
                plt.ylabel('误差(m)')
                #plt.tight_layout()
                plt.legend(loc=1,frameon=False,ncol=4)     
                plt.savefig(savepath + 'prn_{}.png'.format(prn))   
                if display==True:
                    plt.show()
                plt.clf()
            
                    

        if mode=='auto':
            if sv_num <3:
                sub(savepath,sv_num,datas,display)
            else:
                single(savepath,sv_num,datas,display)
        elif mode=='single':
            single(savepath,sv_num,datas,display)
        else:
            sub(savepath,sv_num,datas,display)

    def fixp(self,x,y,z):
        ''' 
        修正坐标到itrf2014坐标系
        '''
        t1=0.060
        t2 = -0.517
        t3= -0.223
        d = -0.011
        r1= 0.0183
        r2= -0.0003
        r3 = 0.0070
        p1 = np.array([[x,y,z]]).T
        p2 = np.array([[t1,t2,t3]]).T
        p3 = np.array([[d,-r3,r2],[r3,d,-r1],[-r3,r1,d]])
        mind = lg.inv(np.ones((3,3))+p3)
        
        r = mind.dot(p1-p2) 
        fx = r.T[0][0]
        fy = r.T[0][1]
        fz = r.T[0][2]

        #fx = (p1 + p2 + p3.dot(p1)).T[0][0]
        #fy = (p1 + p2 + p3.dot(p1)).T[0][1]
        #fz = (p1 + p2 + p3.dot(p1)).T[0][2]
        

        return [fx,fy,fz]


    def run(self,cmpmode='both',display=False):
        '''
        just do it !
        mode:
        1.both: 不外推，精确找都有的时间点计算匹配
        2.sp3: 以sp3的时间为基准，每15分钟计算一次，需要外推星历。
        !!! boy next door warning!!!
        '''        
        self.get_available_data(self.sv_need)
        svs = self.svs
        times = self.times
        nav = self.nav
        sp3 = self.sp3
        t0 = self.t0
        #如果是sp3的话，正常获取数据，但是时间必须要外推。
        
        if cmpmode=='sp3':
            times_temp = []
            for t in times:
                times_temp.append(t)
                times_temp.append(t+np.timedelta64(15,'m'))
                times_temp.append(t+np.timedelta64(30,'m'))
                times_temp.append(t+np.timedelta64(45,'m'))
            times = []
            times = times_temp
             
        # 获取数据
        
        #以卫星为基准
        
        
        sv_data = {}
        
        for sv in svs:
            
            sv_line = {}
            
            xtime = []
            xdata = []
            ydata = []
            zdata = []
            p_igs = []
            p_cal = []
            #p_fix = []  #'转换后的坐标'
            for time in times:
                
                if cmpmode=='sp3':
                    delta = (time - t0) / np.timedelta64('15','m') % 4
                    timet = time - delta * np.timedelta64('15','m')
                    data = nav.sel(sv=sv,time=timet)                    
                else:
                    data = nav.sel(sv=sv,time=time)
                t= int(str(time - t0).split(' ')[0])/1000000000

                #print(data.sv.values,str(data.time.values).split('.')[0]) 
                inf = str(data.sv.values) + '  ' + str(data.time.values).split('.')[0]
                self.logfile.write(inf+ '\r\n')
                self.cal(t,data)
                # 添加数据
                p_cal.append([self.x,self.y,self.z])
                self.logfile.write('计算坐标:{},{},{}\r\n'.format(self.x,self.y,self.z))
                #修正坐标
                '''
                fix = self.fixp(self.x,self.y,self.z)
                p_fix.append(fix)
                self.logfile.write('修正坐标:{},{},{}\r\n'.format(fix[0],fix[1],fix[2]))
                del fix
                '''
                # 比较
                spadata = sp3.sel(sv=sv,time=time)['position'].values
                x = as_num3(spadata[0] * 1000)
                y = as_num3(spadata[1] * 1000)
                z = as_num3(spadata[2] * 1000)
                

                #print('精密坐标:',x,y,z)
                self.logfile.write('精密坐标:{},{},{}\r\n'.format(x,y,z))
                p_igs.append([x,y,z])
                #误差计算
                dx = as_num3(self.x - x)
                dy = as_num3(self.y - y)
                dz = as_num3(self.z - z)
                xlabel  = str(time).split('T')[1].split('.')[0][:-3]
                

                #赋值
                xtime.append(xlabel)
                xdata.append(dx)
                ydata.append(dy)
                zdata.append(dz)
                self.logfile.write('误差(m):{},{},{}'.format(dx,dy,dz))
                self.logfile.write('\r\n\r\n\r\n')
                #print('误差(m):',dx,dy,dz)
                #print()
                #print()

            sv_line['time'] = xtime
            sv_line['x'] = xdata
            sv_line['y'] = ydata
            sv_line['z'] = zdata
            sv_line['p_igs'] = p_igs
            sv_line['p_cal'] = p_cal
            #均值和标准差计算
            '''
            mv 均值 [x,y,z]
            sd 标准差 [x,y,z]
            '''
            mv = []
            sd = []
            mv.append(as_num3(reduce(lambda x,y:x+y , xdata) / len(xdata) ) )
            mv.append(as_num3(reduce(lambda x,y:x+y , ydata) / len(ydata) ) )
            mv.append(as_num3(reduce(lambda x,y:x+y , zdata) / len(zdata) ) )
            sv_line['mv'] = mv
            #先平方来一手
            xdata2 = list(map(lambda x:x*x,xdata))
            ydata2 = list(map(lambda x:x*x,ydata))
            zdata2 = list(map(lambda x:x*x,zdata))
            sd.append(as_num3(np.sqrt(reduce(lambda x,y:x+y , xdata2) / len(xdata2))) )
            sd.append(as_num3(np.sqrt(reduce(lambda x,y:x+y , ydata2) / len(ydata2))) )
            sd.append(as_num3(np.sqrt(reduce(lambda x,y:x+y , zdata2) / len(zdata2))) )
            sv_line['sd'] = sd
            print('{}的{}的误差均值是{}'.format(self.navname,sv,mv))
            print('{}的{}的误差标准差是{}'.format(self.navname,sv,sd))
            #sv_line['p_fix'] = p_fix
            sv_data[sv] = sv_line
        #比较画图
        self.draw(sv_data,mode = self.drawmode,display=display)
        self.logfile.close()
        
        return sv_data
    def get_navname(self):
        return self.navname



def report(datas):
    '''生成相关的表格报告，docx格式，包括:
        1.生成不同观测站之间的x,y,z对比。
        2.生成表格
    '''
    
    final = datas
    #获取数据库中需要的信息
    sp3_list = []
    nav_list = []
    sv_list = []
    for sp3name,data1 in datas.items():
        sp3_list.append(sp3name)
        for navname,navdata in data1.items():
            nav_list.append(navname)
            sv_list.append(set(navdata.keys()))
    #获取都有的卫星
    sv_list = list( reduce(lambda x,y:x&y,sv_list)  )
    #print(sp3_list)
    #print(nav_list)
    #print(sv_list)
    if len(sp3_list) != 1:
        print('多个sp3文件的不同观测站的对比为上线')
        return final
    #1.不同观测站之间的差值。
    # 先来一个遍历比较吧，获取比比较的可能
    num_nav = len(nav_list)
    
    for i in range(num_nav):

        if i+1 != num_nav:

            for j in range(i+1,num_nav,1):
                #此时是观测站序号i和序号j进行对比
                
                cmpdata = {}
                for sv in sv_list:
                    data1 = datas[sp3_list[0]][nav_list[i]][sv]
                    data2 = datas[sp3_list[0]][nav_list[j]][sv]
                    # 对比
                    # 1.获取都有的时间线
                    times1 = data1['time']
                    times2 = data2['time']
                    times = list(set(times1) & set(times2))
                    # 2.根据都有的时间线进行对比
                    
                    chax = []
                    chay = []
                    chaz = []
                    for time in times:
                        index = data1['time'].index(time)
                        chax.append(data1['p_cal'][index][0] - data2['p_cal'][index][0])
                        chay.append(data1['p_cal'][index][1] - data2['p_cal'][index][1])
                        chaz.append(data1['p_cal'][index][2] - data2['p_cal'][index][2])
                    mvx = reduce(lambda x,y:x+y ,chax) / len(chax)
                    mvy = reduce(lambda x,y:x+y ,chay) / len(chay)
                    mvz = reduce(lambda x,y:x+y ,chaz) / len(chaz)
                    mvx = as_num(mvx)
                    mvy = as_num(mvy)
                    mvz = as_num(mvz)
                    chax2 = list(map(lambda x:x*x,chax))
                    chay2 = list(map(lambda x:x*x,chay))
                    chaz2 = list(map(lambda x:x*x,chaz))
                    sdx = np.sqrt(reduce(lambda x,y:x+y , chax2) / len(chax2))
                    sdy = np.sqrt(reduce(lambda x,y:x+y , chay2) / len(chay2))
                    sdz = np.sqrt(reduce(lambda x,y:x+y , chaz2) / len(chaz2))
                    sdx = as_num(sdx)
                    sdy = as_num(sdy)
                    sdz = as_num(sdz)
                    # 输出保存
                    cmpdata[sv] =  {'mv': [mvx,mvy,mvz],'sd':[sdx,sdy,sdz]}
                    print(mvx,mvy,mvz)
                    
                    '''
                    print(nav_list[i],nav_list[j],sv)
                    
                    print(sdx,sdy,sdz)
                    '''
                final[nav_list[i]+'_'+nav_list[j]] = cmpdata
                
    return final


def run(data_path,sv_need,drawmode='auto',cmpmode='sp3',display=False,savepath=r'./数据输出'):
    path_list = os.listdir(data_path)
    nav_list1 = list(filter(lambda x:x.split('.')[-1]=='rnx',path_list))
    sp3_list1 = list(filter(lambda x:x.split('.')[-1]=='sp3',path_list))
    nav_list = list(map(lambda x:data_path+ r'/' + x,nav_list1))
    sp3_list = list(map(lambda x:data_path+ r'/' + x,sp3_list1))
    l = len(sp3_list) * len(nav_list)
    #print('共有',l,'批数据，处理中。。。')
    i =1
    datas = {}
    for sp3 in sp3_list:
        each_sp3 = {}
        
        for nav in nav_list:
            instance = pyec(nav,sp3,sv_need,drawmode,i,l,savepath)
            each_sp3[instance.get_navname()]=  instance.run(cmpmode,display)
            del instance
            i= i +1
        name = str(sp3).split('/')[-1].split('.')[0]
        datas[ name ] = each_sp3

    final = report(datas)
    db = open(savepath+ r'/' + r'db.json',mode='w',encoding='utf-8')
    db.write(json.dumps(final,indent=4,separators=(',',':')))
    db.close()   
    return final

def table_cont(savepath):
    document = Document()
    table = document.add_table(rows=10,cols=14,style='Table Grid')
    document.save(savepath + '/table-1.docx')
    document1 = Document(savepath + '/table-1.docx')
    table = document1.tables[0]
    for row,obj_row in enumerate(table.rows):
        for col,cell in enumerate(obj_row.cells):
            cell.text = cell.text + "%d,%d " % (row,col)
    document1.save(savepath + '/table-2.docx')

def i_am_lazy1(datas,savepath):
    #读取数据
    
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    #第一个表格
    doc.add_heading('表1 计算星历与IGS星历的误差、误差标准差', level=1)
    table1 = doc.add_table(rows=13,cols=5,style='Table Grid')
    table1.cell(0,0).text = '观测站'
    table1.cell(0,1).text = '卫星号'
    table1.cell(0,2).text = '检测量'
    table1.cell(0,3).text = '误差均值'
    table1.cell(0,4).text = '误差标准差'
    #合并观测站和卫星号
    table1.cell(1,0).merge(table1.cell(6,0))
    table1.cell(1,0).text = 'JFNG观测站'
    table1.cell(7,0).merge(table1.cell(12,0))
    table1.cell(7,0).text = 'HKSL观测站'
    table1.cell(1,1).merge(table1.cell(3,1))
    table1.cell(1,1).text = 'G31'
    table1.cell(4,1).merge(table1.cell(6,1))
    table1.cell(4,1).text = 'G32'
    table1.cell(7,1).merge(table1.cell(9,1))
    table1.cell(7,1).text = 'G31'
    table1.cell(10,1).merge(table1.cell(12,1))
    table1.cell(10,1).text = 'G32'
    #表头数据
    def writexyz(table,r,c,l):
        if l=='shu':
            table.cell(r,c).text = 'x'
            table.cell(r+1,c).text = 'y'
            table.cell(r+2,c).text = 'z'
        else:
            table.cell(r,c).text = 'x'
            table.cell(r,c+1).text = 'y'
            table.cell(r,c+2).text = 'z'
    writexyz(table1,1,2,'shu')
    writexyz(table1,4,2,'shu')
    writexyz(table1,7,2,'shu')
    writexyz(table1,10,2,'shu')
    # 写入数据
    def writedata(table,r,c,na,sv,data,typ):
        table.cell(r,c).text = str(data['igs20950'][na][sv][typ][0])
        table.cell(r+1,c).text =str(data['igs20950'][na][sv][typ][1])
        table.cell(r+2,c).text =str(data['igs20950'][na][sv][typ][2])

    writedata(table1,1,3,'JFNG','G31',datas,'mv')
    writedata(table1,4,3,'JFNG','G32',datas,'mv')
    writedata(table1,7,3,'HKSL','G31',datas,'mv')
    writedata(table1,10,3,'HKSL','G32',datas,'mv')

    writedata(table1,1,4,'JFNG','G31',datas,'sd')
    writedata(table1,4,4,'JFNG','G32',datas,'sd')
    writedata(table1,7,4,'HKSL','G31',datas,'sd')
    writedata(table1,10,4,'HKSL','G32',datas,'sd')
    # 表1 完成
   
    # 表2
    doc.add_heading('表2 JFNG与HKSL的坐标差、坐标差标准差', level=1)
    table2 = doc.add_table(rows=7,cols=4,style='Table Grid')
    table2.cell(0,0).text = '卫星号'
    table2.cell(0,1).text = '检测量'
    table2.cell(0,2).text = '坐标差均值'
    table2.cell(0,3).text = '坐标差标准差'
    table2.cell(1,0).merge(table2.cell(3,0))
    table2.cell(1,0).text = 'G31'
    table2.cell(4,0).merge(table2.cell(6,0))
    table2.cell(4,0).text = 'G32'

    writexyz(table2,1,1,'shu')
    writexyz(table2,4,1,'shu')

    def writecmpdata(table,r,c,na,sv,data,typ):
        table.cell(r,c).text = str(data[na][sv][typ][0])
        table.cell(r+1,c).text = str(data[na][sv][typ][1])
        table.cell(r+2,c).text = str(data[na][sv][typ][2])

    writecmpdata(table2,1,2,'JFNG_HKSL','G31',datas,'mv')
    writecmpdata(table2,4,2,'JFNG_HKSL','G32',datas,'mv')

    writecmpdata(table2,1,3,'JFNG_HKSL','G31',datas,'sd')
    writecmpdata(table2,4,3,'JFNG_HKSL','G32',datas,'sd')

   
    
    def draw_table34(name,nav,doc,typ):
        def writexyz(table,r,c,l):
            if l=='shu':
                table.cell(r,c).text = 'x'
                table.cell(r+1,c).text = 'y'
                table.cell(r+2,c).text = 'z'
            else:
                table.cell(r,c).text = 'x'
                table.cell(r,c+1).text = 'y'
                table.cell(r,c+2).text = 'z'
        
        doc.add_heading(name, level=1)
        if typ=='shu':
            
            table3 = doc.add_table(rows=14,cols=10,style='Table Grid')
            table3.cell(0,0).merge(table3.cell(1,0))
            table3.cell(0,0).text = '时间'
            table3.cell(0,1).merge(table3.cell(0,3))
            table3.cell(0,1).text = 'IGS坐标'
            table3.cell(0,4).merge(table3.cell(0,6))
            table3.cell(0,4).text = '计算坐标'
            table3.cell(0,7).merge(table3.cell(0,9))
            table3.cell(0,7).text = '误差(m)'
            #写时间
            times = datas['igs20950']['JFNG'][nav]['time']
            for i in range(1,13):
                table3.cell(i+1,0).text = times[i-1]
            #写xyz
            writexyz(table3,1,1,'heng')
            writexyz(table3,1,4,'heng')
            writexyz(table3,1,7,'heng')
            #写实际数据
            pigs = datas['igs20950']['JFNG'][nav]['p_igs']
            pcal = datas['igs20950']['JFNG'][nav]['p_cal']
            
            for i in range(1,13):
                table3.cell(i+1,1).text = str(pigs[i-1][0])
                table3.cell(i+1,2).text = str(pigs[i-1][1])
                table3.cell(i+1,3).text = str(pigs[i-1][2])
                table3.cell(i+1,4).text = str(pcal[i-1][0])
                table3.cell(i+1,5).text = str(pcal[i-1][1])
                table3.cell(i+1,6).text = str(pcal[i-1][2])
                table3.cell(i+1,7).text = str(datas['igs20950']['JFNG'][nav]['x'][i-1])
                table3.cell(i+1,8).text = str(datas['igs20950']['JFNG'][nav]['y'][i-1])
                table3.cell(i+1,9).text = str(datas['igs20950']['JFNG'][nav]['z'][i-1])

        else:
            table3 = doc.add_table(rows=10,cols=14,style='Table Grid')
            table3.cell(0,0).merge(table3.cell(0,1))
            table3.cell(0,0).text = '时间'
            table3.cell(1,0).merge(table3.cell(3,0))
            table3.cell(1,0).text = 'IGS坐标'
            table3.cell(4,0).merge(table3.cell(6,0))
            table3.cell(4,0).text = '计算坐标'
            table3.cell(7,0).merge(table3.cell(9,0))
            table3.cell(7,0).text = '误差(m)'
            times = datas['igs20950']['JFNG']['G31']['time']
            for i in range(1,13):
                table3.cell(0,i+1).text = times[i-1]
            #写xyz
            writexyz(table3,1,1,'shu')
            writexyz(table3,4,1,'shu')
            writexyz(table3,7,1,'shu')
            #写实际数据
            pigs = datas['igs20950']['JFNG']['G31']['p_igs']
            pcal = datas['igs20950']['JFNG']['G31']['p_cal']
            
            for i in range(1,13):
                table3.cell(1,i+1).text = str(pigs[i-1][0])
                table3.cell(2,i+1).text = str(pigs[i-1][1])
                table3.cell(3,i+1).text = str(pigs[i-1][2])
                table3.cell(4,i+1).text = str(pcal[i-1][0])
                table3.cell(5,i+1).text = str(pcal[i-1][1])
                table3.cell(6,i+1).text = str(pcal[i-1][2])
                table3.cell(7,i+1).text = str(datas['igs20950']['JFNG']['G31']['x'][i-1])
                table3.cell(8,i+1).text = str(datas['igs20950']['JFNG']['G31']['y'][i-1])
                table3.cell(9,i+1).text = str(datas['igs20950']['JFNG']['G31']['z'][i-1])

        
    #表31
    #表4
    draw_table34('表3 G31的坐标对比(时间竖向)','G31',doc,'shu')
    draw_table34('表4 G32的坐标对比(时间竖向)','G32',doc,'shu')
    draw_table34('表5 G31的坐标对比(时间横向)','G31',doc,'heng')
    draw_table34('表6 G32的坐标对比(时间横向)','G32',doc,'heng')


    doc.save(savepath + r'/数据表格.docx')
    
    pass     
    
